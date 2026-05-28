import io
from datetime import datetime, timedelta
from types import SimpleNamespace

import pytest
from docx import Document

from backend.application.task_file_extractor_service import TaskFileExtractorService
from backend.application.task_schedule_block_service import TaskScheduleBlockService
from backend.application.task_scheduler_service import TaskSchedulerService


class FakeQuery:
    def __init__(self, items):
        self.items = list(items)

    def filter(self, *args, **kwargs):
        return self

    def all(self):
        return list(self.items)

    def first(self):
        return self.items[0] if self.items else None


class FakeDB:
    def __init__(self, items=None):
        self.items = items or []
        self.added = []
        self.deleted = []
        self.flushed = False

    def query(self, model):
        return FakeQuery(self.items)

    def add(self, item):
        self.added.append(item)

    def delete(self, item):
        self.deleted.append(item)

    def flush(self):
        self.flushed = True


def build_task_docx_bytes():
    document = Document()
    document.add_paragraph("Лабораторна робота №1")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Завдання"
    table.cell(0, 1).text = "Розробити застосунок"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_task_file_extractor_reads_txt_and_cleans_noise_lines():
    service = TaskFileExtractorService()
    result = service.extract_text("task.txt", "1\nЛабораторна   робота\n\n№ 1".encode("utf-8"))

    assert "Лабораторна робота" in result
    assert "№1" in result
    assert "\n1\n" not in result


def test_task_file_extractor_reads_docx_paragraphs_and_tables():
    service = TaskFileExtractorService()
    result = service.extract_text("task.docx", build_task_docx_bytes())

    assert "Лабораторна робота" in result
    assert "Завдання | Розробити застосунок" in result


def test_task_file_extractor_rejects_unsupported_extension():
    service = TaskFileExtractorService()

    with pytest.raises(ValueError):
        service.extract_text("archive.zip", b"data")


def test_task_schedule_block_service_recreates_block_and_deletes_old_blocks():
    old_block = SimpleNamespace(id=1)
    db = FakeDB(items=[old_block])
    service = TaskScheduleBlockService()
    task = SimpleNamespace(id=10, estimated_duration_hours=10)
    deadline = datetime(2026, 5, 28, 20, 0)

    block = service.recreate_block_for_task(
        db=db,
        user_id=5,
        task=task,
        deadline=deadline,
        confidence_score=0.66,
    )

    assert old_block in db.deleted
    assert block.start_time == deadline - timedelta(hours=3)
    assert block.end_time == deadline
    assert block.generated_by_ai is True
    assert block.source == "ml_deadline_planner"
    assert block.confidence_score == 0.66
    assert db.flushed is True


def test_task_scheduler_service_moves_block_after_overlapping_event():
    service = TaskSchedulerService()
    event = SimpleNamespace(
        start_time=datetime(2026, 5, 28, 9, 0),
        end_time=datetime(2026, 5, 28, 11, 0),
    )
    db = FakeDB(items=[event])
    task = SimpleNamespace(id=77, estimated_duration_hours=2)
    deadline = datetime(2026, 5, 28, 10, 0)

    block = service.generate_task_blocks(db=db, user_id=1, task=task, deadline=deadline)

    assert block.start_time == datetime(2026, 5, 28, 11, 30)
    assert block.end_time == datetime(2026, 5, 28, 13, 30)
    assert block.source == "ml_scheduler"
    assert db.added == [block]
