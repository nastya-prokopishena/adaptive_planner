import io

import fitz
import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image

from backend.application.schedule_file_extractor_service import ScheduleFileExtractorService
from backend.infrastructure.file_extractors.docx_extractor import DocxExtractor
from backend.infrastructure.file_extractors.excel_extractor import ExcelExtractor
from backend.infrastructure.file_extractors.image_extractor import ImageExtractor
from backend.infrastructure.file_extractors.pdf_extractor import PdfExtractor


def build_docx_bytes():
    document = Document()
    document.add_paragraph("Лекція 1. Архітектура програмного забезпечення")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "День"
    table.cell(0, 1).text = "Пара"
    table.cell(1, 0).text = "Понеділок"
    table.cell(1, 1).text = "ООП"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_xlsx_bytes():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Schedule"
    sheet.append(["День", "Час", "Предмет"])
    sheet.append(["Понеділок", "08:30", "Бази даних"])
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def build_png_bytes():
    image = Image.new("RGB", (20, 20), color="white")
    output = io.BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


def build_pdf_bytes():
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Розклад групи ФеП-42")
    pdf_bytes = document.tobytes()
    document.close()
    return pdf_bytes


def test_docx_extractor_reads_paragraphs_and_tables():
    result = DocxExtractor().extract("schedule.docx", build_docx_bytes())

    assert result["extension"] == "docx"
    assert "Архітектура" in result["text_context"]
    assert result["tables"]
    assert result["debug"]["extractor"] == "docx_text_and_tables_strategy"


def test_excel_extractor_reads_xlsx_rows():
    result = ExcelExtractor("xlsx").extract("schedule.xlsx", build_xlsx_bytes())

    assert result["extension"] == "xlsx"
    assert "Бази даних" in result["text_context"]
    assert result["tables"][0]["name"] == "Schedule"


def test_image_extractor_converts_image_to_base64_page():
    result = ImageExtractor("png").extract("photo.png", build_png_bytes())

    assert result["extension"] == "png"
    assert result["pages"][0]["full_image"]["mime_type"] == "image/png"
    assert result["pages"][0]["full_image"]["base64"]


def test_pdf_extractor_reads_text_and_generates_page_image():
    result = PdfExtractor().extract("schedule.pdf", build_pdf_bytes())

    assert result["extension"] == "pdf"
    assert result["text_context"]
    assert result["pages"][0]["full_image"]["base64"]


def test_schedule_file_extractor_service_adds_target_group_debug():
    service = ScheduleFileExtractorService()

    result = service.extract(
        filename="schedule.txt",
        file_bytes="Понеділок | 08:30 | ООП".encode("utf-8"),
        group_name="ФеП-42",
    )

    assert result["extension"] == "txt"
    assert result["debug"]["target_group"] == "ФеП-42"


def test_schedule_file_extractor_service_truncates_manual_text():
    service = ScheduleFileExtractorService()

    result = service.extract_text_input("a" * 170_000)

    assert result["debug"]["text_chars"] <= 160_025
    assert "TEXT TRUNCATED" in result["text_context"]
