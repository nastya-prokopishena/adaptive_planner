import io

from docx import Document

from backend.domain.interfaces.file_extractor import FileExtractor
from backend.infrastructure.file_extractors.base_extractor import BaseExtractor


class DocxExtractor(BaseExtractor, FileExtractor):
    def extract(self, filename: str, file_bytes: bytes):
        document = Document(io.BytesIO(file_bytes))

        parts = []
        tables = []

        for paragraph in document.paragraphs:
            text = self._clean_text(paragraph.text)

            if text:
                parts.append(text)

        for table_index, table in enumerate(document.tables, start=1):
            parts.append(f"\n--- DOCX TABLE {table_index} ---")

            table_rows = []

            for row in table.rows:
                cells = [self._clean_text(cell.text) for cell in row.cells]

                if any(cells):
                    table_rows.append(cells)
                    parts.append(" | ".join(cells))

            tables.append(
                {
                    "name": f"DOCX TABLE {table_index}",
                    "rows": table_rows,
                }
            )

        return self._text_result(
            filename=filename,
            extension="docx",
            text="\n".join(parts),
            extractor_name="docx_text_and_tables_strategy",
            tables=tables,
        )
