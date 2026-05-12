import base64
import io
import json
import mimetypes
import os
import re
from typing import Any

import fitz
import pandas as pd
from openai import OpenAI


class AIScheduleParserService:
    IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp"}
    PDF_EXTENSIONS = {"pdf"}
    EXCEL_EXTENSIONS = {"xlsx", "xls"}
    DOCX_EXTENSIONS = {"docx"}
    TEXT_EXTENSIONS = {"txt", "csv"}

    DEFAULT_PAIR_TIMES = {
        1: ("08:30", "09:50"),
        2: ("10:10", "11:30"),
        3: ("11:50", "13:10"),
        4: ("13:30", "14:50"),
        5: ("15:05", "16:25"),
        6: ("16:40", "18:00"),
        7: ("18:10", "19:30"),
        8: ("19:40", "21:00"),
    }

    def __init__(self):
        self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.model = os.getenv("OPENAI_SCHEDULE_MODEL", "gpt-4.1")

    def parse_file(
        self,
        filename: str,
        file_bytes: bytes,
        group_name: str = "",
        subgroup: str = "",
    ) -> list[dict[str, Any]]:
        extension = self._get_extension(filename)

        if extension in self.IMAGE_EXTENSIONS:
            return self._parse_image(
                filename=filename,
                file_bytes=file_bytes,
                group_name=group_name,
                subgroup=subgroup,
            )

        if extension in self.PDF_EXTENSIONS:
            return self._parse_pdf_as_images(
                filename=filename,
                file_bytes=file_bytes,
                group_name=group_name,
                subgroup=subgroup,
            )

        if extension in self.EXCEL_EXTENSIONS:
            extracted_text = self._extract_excel_text(file_bytes, extension)
            return self.parse_text(
                text=extracted_text,
                group_name=group_name,
                subgroup=subgroup,
                source_hint=f"Excel file: {filename}",
            )

        if extension in self.DOCX_EXTENSIONS:
            extracted_text = self._extract_docx_text(file_bytes)
            return self.parse_text(
                text=extracted_text,
                group_name=group_name,
                subgroup=subgroup,
                source_hint=f"DOCX file: {filename}",
            )

        if extension in self.TEXT_EXTENSIONS:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")
            return self.parse_text(
                text=extracted_text,
                group_name=group_name,
                subgroup=subgroup,
                source_hint=f"Text file: {filename}",
            )

        raise ValueError(
            "Непідтримуваний формат файлу. Доступні формати: PDF, XLSX, XLS, CSV, DOCX, TXT, JPG, PNG, WEBP."
        )

    def parse_text(
        self,
        text: str,
        group_name: str = "",
        subgroup: str = "",
        source_hint: str = "raw text",
    ) -> list[dict[str, Any]]:
        if not text or not text.strip():
            return []

        prompt = self._build_prompt(
            group_name=group_name,
            subgroup=subgroup,
            source_hint=source_hint,
        )

        response = self.client.responses.create(
            model=self.model,
            input=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_text",
                            "text": f"{prompt}\n\nТЕКСТ РОЗКЛАДУ:\n{text}",
                        }
                    ],
                }
            ],
        )

        return self._parse_ai_json_response(response.output_text)

    def _parse_image(
        self,
        filename: str,
        file_bytes: bytes,
        group_name: str = "",
        subgroup: str = "",
    ) -> list[dict[str, Any]]:
        mime_type = mimetypes.guess_type(filename)[0] or "image/png"
        image_base64 = base64.b64encode(file_bytes).decode("utf-8")
        image_url = f"data:{mime_type};base64,{image_base64}"

        prompt = self._build_prompt(
            group_name=group_name,
            subgroup=subgroup,
            source_hint=f"Image schedule file: {filename}",
        )

        response = self.client.responses.create(
            model=self.model,
            input=[
                {
                    "role": "user",
                    "content": [
                        {"type": "input_text", "text": prompt},
                        {"type": "input_image", "image_url": image_url},
                    ],
                }
            ],
        )

        return self._parse_ai_json_response(response.output_text)

    def _parse_pdf_as_images(
        self,
        filename: str,
        file_bytes: bytes,
        group_name: str = "",
        subgroup: str = "",
    ) -> list[dict[str, Any]]:
        page_images = self._render_pdf_pages_to_data_urls(file_bytes)

        if not page_images:
            return []

        prompt = self._build_prompt(
            group_name=group_name,
            subgroup=subgroup,
            source_hint=f"PDF rendered as page images: {filename}",
        )

        content: list[dict[str, Any]] = [
            {
                "type": "input_text",
                "text": prompt,
            }
        ]

        for index, image_url in enumerate(page_images, start=1):
            content.append(
                {
                    "type": "input_text",
                    "text": (
                        f"Сторінка PDF №{index}. "
                        f"Проаналізуй її саме візуально як таблицю: колонки груп, об’єднані клітинки, "
                        f"потоки, підгрупи, дні, час і тижні."
                    ),
                }
            )
            content.append(
                {
                    "type": "input_image",
                    "image_url": image_url,
                }
            )

        response = self.client.responses.create(
            model=self.model,
            input=[
                {
                    "role": "user",
                    "content": content,
                }
            ],
        )

        return self._parse_ai_json_response(response.output_text)

    def _render_pdf_pages_to_data_urls(self, file_bytes: bytes) -> list[str]:
        document = fitz.open(stream=file_bytes, filetype="pdf")
        images: list[str] = []

        try:
            for page in document:
                matrix = fitz.Matrix(2.5, 2.5)
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)

                image_bytes = pixmap.tobytes("png")
                image_base64 = base64.b64encode(image_bytes).decode("utf-8")
                images.append(f"data:image/png;base64,{image_base64}")
        finally:
            document.close()

        return images

    def _extract_excel_text(self, file_bytes: bytes, extension: str) -> str:
        engine = "openpyxl" if extension == "xlsx" else "xlrd"
        result_parts: list[str] = []

        excel = pd.ExcelFile(io.BytesIO(file_bytes), engine=engine)

        for sheet_name in excel.sheet_names:
            df = pd.read_excel(
                io.BytesIO(file_bytes),
                sheet_name=sheet_name,
                header=None,
                dtype=str,
                engine=engine,
            )

            result_parts.append(f"\n=== SHEET: {sheet_name} ===")

            for _, row in df.iterrows():
                values = []

                for value in row.tolist():
                    cleaned = self._clean_cell(value)

                    if cleaned:
                        values.append(cleaned)

                if values:
                    result_parts.append(" | ".join(values))

        return "\n".join(result_parts)

    def _extract_docx_text(self, file_bytes: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Для DOCX потрібно встановити python-docx.") from exc

        document = Document(io.BytesIO(file_bytes))
        result_parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                result_parts.append(text)

        for table in document.tables:
            result_parts.append("\n=== TABLE ===")

            for row in table.rows:
                values = []

                for cell in row.cells:
                    cleaned = self._clean_cell(cell.text)

                    if cleaned:
                        values.append(cleaned)

                if values:
                    result_parts.append(" | ".join(values))

        return "\n".join(result_parts)

    def _build_prompt(
        self,
        group_name: str = "",
        subgroup: str = "",
        source_hint: str = "",
    ) -> str:
        return f"""
Ти модуль для точного зчитування університетського розкладу з таблиць, PDF, фото, Excel, DOCX і тексту.

Потрібно повернути ТІЛЬКИ JSON без markdown, без пояснень, без тексту навколо.

Користувач шукає розклад тільки для:
group_name = "{group_name}"
subgroup = "{subgroup}"

Твоє головне завдання:
1. Не вигадувати пари.
2. Не приписувати чужі групи до "{group_name}".
3. Додати тільки ті заняття, які реально належать групі "{group_name}".
4. Правильно відрізняти групи від підгруп.

КРИТИЧНІ ПРАВИЛА ДЛЯ ГРУП:

1. Групи — це назви типу:
   ФЕП-42с, ФЕП-41с, ФЕП-43с, ТЯ-2304, КН-21 тощо.

2. Підгрупи — це тільки:
   "підгр. 1", "підгр. 2", "підгрупа 1", "підгрупа 2".

3. "гр. 1", "група 1", "Зб. група", "Потік" НЕ є підгрупою.

4. Для кожної події обов’язково визнач:
   - source_group_text — яку групу/групи реально видно біля заняття;
   - group_relation — як саме заняття пов’язане з group_name "{group_name}".

Дозволені значення group_relation:
- "exact_column" — заняття стоїть прямо в колонці "{group_name}";
- "merged_over_target_group" — заняття в об’єднаній/розтягнутій клітинці, яка перетинає колонку "{group_name}";
- "other_group" — заняття належить іншій групі;
- "unknown" — неможливо довести, що заняття належить "{group_name}".

ВАЖЛИВО:
Якщо group_relation не "exact_column" і не "merged_over_target_group", backend потім видалить цю подію.
Тому не став "exact_column" або "merged_over_target_group", якщо це не видно з таблиці.

ПРАВИЛА ДЛЯ ОБ’ЄДНАНИХ КЛІТИНОК І ПОТОКІВ:

5. Якщо лекція або потік розтягнуті на кілька колонок груп і ця велика клітинка перетинає колонку "{group_name}", цю пару треба додати.

6. Якщо велика клітинка з лекцією розтягнута на ФЕП-41с, ФЕП-42с, ФЕП-43с, і користувач просить ФЕП-42с — додай її з:
   group_relation = "merged_over_target_group"
   source_group_text = "ФЕП-41с, ФЕП-42с, ФЕП-43с" або "потік"

7. Якщо заняття стоїть тільки в колонці іншої групи і не перетинає колонку "{group_name}" — не додавай його або постав group_relation = "other_group".

ПРАВИЛА ДЛЯ ПІДГРУП:

8. Якщо subgroup = "{subgroup}", включай:
   - заняття з явним "підгр. {subgroup}";
   - заняття без підгрупи, якщо це лекція / потік / заняття для всієї групи.

9. Якщо subgroup = "{subgroup}", НЕ включай:
   - "підгр. 2", якщо користувач просить "1";
   - "підгр. 1", якщо користувач просить "2";
   - лабораторні/практичні/семінари без підгрупи, якщо не видно, що вони саме для цієї підгрупи.

10. Якщо лабораторна або практична має іншу підгрупу — не додавай її.

11. Для кожної події поверни subgroup_source:
   - "explicit" — якщо підгрупа явно написана як "підгр. 1";
   - "none" — якщо підгрупи немає, бо це лекція/потік/вся група;
   - "unknown" — якщо не зрозуміло.

ПРАВИЛА ДЛЯ ТИЖНІВ:

12. Якщо написано "чис.", "чисельник", "н/пар", "непарні" — week_pattern = "odd".
13. Якщо написано "знам.", "знаменник", "пар.", "парні" — week_pattern = "even".
14. Якщо в один день і в один час для тієї самої групи стоять дві різні події, вони не можуть бути щотижня одночасно.
    Якщо явного чис./знам. немає, постав:
    - першій week_pattern = "odd";
    - другій week_pattern = "even";
    - needs_review = true;
    - confidence не вище 0.8.

ПРАВИЛА ДЛЯ ЧАСУ:

15. Якщо вказаний точний час — використовуй його.
16. Якщо є номер пари, але немає часу — використовуй:
   1 пара = 08:30-09:50
   2 пара = 10:10-11:30
   3 пара = 11:50-13:10
   4 пара = 13:30-14:50
   5 пара = 15:05-16:25
   6 пара = 16:40-18:00
   7 пара = 18:10-19:30
   8 пара = 19:40-21:00

ПРАВИЛА ДЛЯ ПОЛІВ:

17. subject — тільки назва предмета. Без викладача, аудиторії, групи, підгрупи, типу заняття, часу.
18. event_type:
   - lecture
   - laboratory
   - practice
   - seminar
   - consultation
   - exam
   - credit
   - class

19. day_of_week:
   - MO
   - TU
   - WE
   - TH
   - FR
   - SA
   - SU

20. Якщо викладача немає — teacher = "".
21. Якщо аудиторії немає — room = "".
22. Якщо підгрупи немає, але це лекція / потік / вся група — subgroup = "" і subgroup_source = "none".

ПОВЕРНИ JSON СТРОГО ТАКОЇ СТРУКТУРИ:

{{
  "events": [
    {{
      "subject": "",
      "event_type": "class",
      "day_of_week": "",
      "start_time": "",
      "end_time": "",
      "teacher": "",
      "room": "",
      "group_name": "{group_name}",
      "source_group_text": "",
      "group_relation": "unknown",
      "subgroup": "",
      "subgroup_source": "unknown",
      "week_pattern": "weekly",
      "confidence": 0.0,
      "needs_review": false
    }}
  ]
}}

НЕ ДОДАВАЙ ПАРИ, ЯКЩО:
- вони явно належать іншій групі;
- вони явно належать іншій підгрупі;
- ти не бачиш, що заняття стоїть у колонці "{group_name}" або merged-коміркою перетинає її.

Контекст джерела: {source_hint}
""".strip()

    def _parse_ai_json_response(self, output_text: str) -> list[dict[str, Any]]:
        if not output_text:
            return []

        cleaned = output_text.strip()
        cleaned = cleaned.replace("```json", "")
        cleaned = cleaned.replace("```", "")
        cleaned = cleaned.strip()

        json_text = self._extract_json_object(cleaned)

        try:
            parsed = json.loads(json_text)
        except json.JSONDecodeError as exc:
            raise ValueError(f"AI повернув невалідний JSON: {cleaned}") from exc

        events = parsed.get("events", [])

        if not isinstance(events, list):
            return []

        normalized_events = []

        for event in events:
            if isinstance(event, dict):
                normalized_events.append(self._normalize_event(event))

        return normalized_events

    def _normalize_event(self, event: dict[str, Any]) -> dict[str, Any]:
        subject = str(event.get("subject") or "").strip()
        event_type = str(event.get("event_type") or "class").strip().lower()
        day_of_week = str(event.get("day_of_week") or "").strip().upper()
        start_time = self._normalize_time(event.get("start_time"))
        end_time = self._normalize_time(event.get("end_time"))
        teacher = str(event.get("teacher") or "").strip()
        room = str(event.get("room") or "").strip()
        group_name = str(event.get("group_name") or "").strip()
        source_group_text = str(event.get("source_group_text") or "").strip()
        group_relation = str(event.get("group_relation") or "unknown").strip().lower()
        subgroup = self._normalize_subgroup(event.get("subgroup"))
        subgroup_source = str(event.get("subgroup_source") or "unknown").strip().lower()
        week_pattern = str(event.get("week_pattern") or "weekly").strip().lower()

        confidence = event.get("confidence", 0.0)

        try:
            confidence = float(confidence)
        except Exception:
            confidence = 0.0

        confidence = max(0.0, min(confidence, 1.0))
        needs_review = bool(event.get("needs_review", False))

        allowed_types = {
            "lecture",
            "laboratory",
            "practice",
            "seminar",
            "consultation",
            "exam",
            "credit",
            "class",
        }

        allowed_days = {"MO", "TU", "WE", "TH", "FR", "SA", "SU"}

        allowed_relations = {
            "exact_column",
            "merged_over_target_group",
            "other_group",
            "unknown",
        }

        allowed_subgroup_sources = {
            "explicit",
            "none",
            "unknown",
        }

        allowed_week_patterns = {
            "weekly",
            "odd",
            "even",
        }

        if event_type not in allowed_types:
            event_type = "class"

        if day_of_week not in allowed_days:
            day_of_week = ""
            needs_review = True
            confidence = min(confidence, 0.6)

        if group_relation not in allowed_relations:
            group_relation = "unknown"

        if subgroup_source not in allowed_subgroup_sources:
            subgroup_source = "unknown"

        if week_pattern not in allowed_week_patterns:
            week_pattern = "weekly"

        if not subject:
            needs_review = True
            confidence = min(confidence, 0.5)

        if not start_time or not end_time:
            needs_review = True
            confidence = min(confidence, 0.75)

        return {
            "subject": subject,
            "event_type": event_type,
            "day_of_week": day_of_week,
            "start_time": start_time,
            "end_time": end_time,
            "teacher": teacher,
            "room": room,
            "group_name": group_name,
            "source_group_text": source_group_text,
            "group_relation": group_relation,
            "subgroup": subgroup,
            "subgroup_source": subgroup_source,
            "week_pattern": week_pattern,
            "confidence": round(confidence, 2),
            "needs_review": needs_review,
        }

    def _extract_json_object(self, text: str) -> str:
        match = re.search(r"\{.*\}", text, flags=re.DOTALL)

        if not match:
            raise ValueError(f"У відповіді AI немає JSON-об'єкта: {text}")

        return match.group(0)

    def _normalize_time(self, value: Any) -> str:
        text = str(value or "").strip().replace(".", ":")

        if not text:
            return ""

        match = re.fullmatch(r"(\d{1,2}):(\d{2})", text)

        if not match:
            return ""

        hours = int(match.group(1))
        minutes = int(match.group(2))

        if hours > 23 or minutes > 59:
            return ""

        return f"{hours:02d}:{minutes:02d}"

    def _normalize_subgroup(self, value: Any) -> str:
        text = str(value or "").lower().strip()

        if not text:
            return ""

        text = text.replace("підгр.", "")
        text = text.replace("підгр", "")
        text = text.replace("підгрупа", "")
        text = text.replace(".", "")
        text = text.replace(" ", "")

        match = re.search(r"\d+", text)

        if match:
            return match.group(0)

        return ""

    def _get_extension(self, filename: str) -> str:
        if "." not in filename:
            return ""

        return filename.rsplit(".", 1)[-1].lower().strip()

    def _clean_cell(self, value: Any) -> str:
        if value is None:
            return ""

        try:
            if pd.isna(value):
                return ""
        except Exception:
            pass

        text = str(value)
        text = text.replace("\r", " ")
        text = text.replace("\n", " ")
        text = re.sub(r"\s+", " ", text)

        return text.strip()