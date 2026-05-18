import base64
import csv
import io
import os
import re
from typing import Any

import fitz
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from PIL import Image


class ScheduleFileExtractorService:
    PDF_EXTENSIONS = {"pdf"}
    IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp"}
    DOCX_EXTENSIONS = {"docx"}
    EXCEL_EXTENSIONS = {"xlsx", "xls"}
    TEXT_EXTENSIONS = {"txt", "csv"}

    MAX_TEXT_CHARS = 160_000
    MAX_PDF_PAGES = 20
    PDF_ZOOM = 3
    MAX_IMAGE_SIDE = 2800

    def extract(self, filename: str, file_bytes: bytes, group_name: str = "") -> dict[str, Any]:
        extension = self._get_extension(filename)

        if extension in self.PDF_EXTENSIONS:
            return self._extract_pdf(filename, file_bytes)

        if extension in self.IMAGE_EXTENSIONS:
            return self._extract_image(filename, file_bytes, extension)

        if extension in self.DOCX_EXTENSIONS:
            return self._extract_docx(filename, file_bytes)

        if extension in self.EXCEL_EXTENSIONS:
            return self._extract_excel(filename, file_bytes, extension)

        if extension in self.TEXT_EXTENSIONS:
            return self._extract_text_or_csv(filename, file_bytes, extension)

        raise ValueError(
            f"Формат .{extension} не підтримується. "
            "Підтримуються PDF, фото, Excel, DOCX, TXT, CSV."
        )

    def extract_text_input(self, raw_text: str) -> dict[str, Any]:
        text = self._truncate_text(raw_text or "")

        return {
            "filename": "manual_text",
            "extension": "txt",
            "text_context": text,
            "pages": [],
            "tables": [],
            "debug": {
                "extractor": "manual_text",
                "text_chars": len(text),
            },
        }

    def _extract_pdf(self, filename: str, file_bytes: bytes) -> dict[str, Any]:
        document = fitz.open(stream=file_bytes, filetype="pdf")

        pages = []
        text_parts = []

        try:
            page_count = min(len(document), self.MAX_PDF_PAGES)

            for page_index in range(page_count):
                page = document[page_index]
                page_number = page_index + 1

                page_text = self._clean_text(page.get_text("text") or "")
                text_parts.append(f"\n\n--- PDF PAGE {page_number} TEXT ---\n{page_text}")

                pixmap = page.get_pixmap(
                    matrix=fitz.Matrix(self.PDF_ZOOM, self.PDF_ZOOM),
                    alpha=False,
                )

                image = Image.open(io.BytesIO(pixmap.tobytes("png"))).convert("RGB")
                image.thumbnail((self.MAX_IMAGE_SIDE, self.MAX_IMAGE_SIDE))

                output = io.BytesIO()
                image.save(output, format="PNG")

                pages.append(
                    {
                        "page": page_number,
                        "page_text": page_text,
                        "full_image": {
                            "filename": f"{filename}_page_{page_number}.png",
                            "mime_type": "image/png",
                            "base64": base64.b64encode(output.getvalue()).decode("utf-8"),
                        },
                    }
                )

        finally:
            document.close()

        text_context = self._truncate_text("\n".join(text_parts))

        return {
            "filename": filename,
            "extension": "pdf",
            "text_context": text_context,
            "pages": pages,
            "tables": [],
            "debug": {
                "extractor": "pdf_visual_and_text",
                "pages": len(pages),
                "text_chars": len(text_context),
                "size_bytes": len(file_bytes),
            },
        }

    def _extract_image(self, filename: str, file_bytes: bytes, extension: str) -> dict[str, Any]:
        image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        image.thumbnail((self.MAX_IMAGE_SIDE, self.MAX_IMAGE_SIDE))

        output = io.BytesIO()
        image.save(output, format="PNG")

        return {
            "filename": filename,
            "extension": extension,
            "text_context": "",
            "pages": [
                {
                    "page": 1,
                    "page_text": "",
                    "full_image": {
                        "filename": filename,
                        "mime_type": "image/png",
                        "base64": base64.b64encode(output.getvalue()).decode("utf-8"),
                    },
                }
            ],
            "tables": [],
            "debug": {
                "extractor": "image_visual",
                "size_bytes": len(file_bytes),
            },
        }

    def _extract_docx(self, filename: str, file_bytes: bytes) -> dict[str, Any]:
        document = Document(io.BytesIO(file_bytes))
        parts = []
        tables = []

        for paragraph in document.paragraphs:
            text = self._clean_text(paragraph.text)
            if text:
                parts.append(text)

        for table_index, table in enumerate(document.tables, start=1):
            parts.append(f"\n--- DOCX TABLE {table_index} ---")

            table_rows = []

            for row in table.rows:
                cells = [self._clean_text(cell.text) for cell in row.cells]
                if any(cells):
                    table_rows.append(cells)
                    parts.append(" | ".join(cells))

            tables.append(
                {
                    "name": f"DOCX TABLE {table_index}",
                    "rows": table_rows,
                }
            )

        return self._text_result(
            filename=filename,
            extension="docx",
            text="\n".join(parts),
            extractor_name="docx_text_and_tables",
            tables=tables,
        )

    def _extract_excel(self, filename: str, file_bytes: bytes, extension: str) -> dict[str, Any]:
        parts = []
        tables = []

        if extension == "xlsx":
            workbook = load_workbook(io.BytesIO(file_bytes), data_only=True)

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                parts.append(f"\n--- EXCEL SHEET: {sheet_name} ---")

                sheet_rows = []

                for row in sheet.iter_rows(values_only=True):
                    cells = [self._clean_text(cell) for cell in row]
                    if any(cells):
                        sheet_rows.append(cells)
                        parts.append(" | ".join(cells))

                tables.append(
                    {
                        "name": sheet_name,
                        "rows": sheet_rows,
                    }
                )

        else:
            excel_file = pd.ExcelFile(io.BytesIO(file_bytes), engine="xlrd")

            for sheet_name in excel_file.sheet_names:
                dataframe = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
                parts.append(f"\n--- EXCEL SHEET: {sheet_name} ---")

                sheet_rows = []

                for _, row in dataframe.iterrows():
                    cells = [self._clean_text(cell) for cell in row.tolist()]
                    if any(cells):
                        sheet_rows.append(cells)
                        parts.append(" | ".join(cells))

                tables.append(
                    {
                        "name": sheet_name,
                        "rows": sheet_rows,
                    }
                )

        return self._text_result(
            filename=filename,
            extension=extension,
            text="\n".join(parts),
            extractor_name=f"{extension}_tables",
            tables=tables,
        )

    def _extract_text_or_csv(self, filename: str, file_bytes: bytes, extension: str) -> dict[str, Any]:
        decoded = self._decode_bytes(file_bytes)

        if extension == "csv":
            reader = csv.reader(io.StringIO(decoded))
            lines = []

            for row in reader:
                lines.append(" | ".join(self._clean_text(cell) for cell in row))

            text = "\n".join(lines)
        else:
            text = decoded

        return self._text_result(filename, extension, text, extension, tables=[])

    def _text_result(
        self,
        filename: str,
        extension: str,
        text: str,
        extractor_name: str,
        tables: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        text = self._truncate_text(text)

        return {
            "filename": filename,
            "extension": extension,
            "text_context": text,
            "pages": [],
            "tables": tables or [],
            "debug": {
                "extractor": extractor_name,
                "text_chars": len(text),
            },
        }

    def _decode_bytes(self, file_bytes: bytes) -> str:
        for encoding in ["utf-8-sig", "utf-8", "cp1251", "latin-1"]:
            try:
                return file_bytes.decode(encoding)
            except Exception:
                continue

        return file_bytes.decode("utf-8", errors="ignore")

    def _truncate_text(self, text: str) -> str:
        text = str(text or "")

        if len(text) <= self.MAX_TEXT_CHARS:
            return text

        return text[: self.MAX_TEXT_CHARS] + "\n\n--- TEXT TRUNCATED ---"

    def _clean_text(self, value: Any) -> str:
        if value is None:
            return ""

        try:
            if pd.isna(value):
                return ""
        except Exception:
            pass

        text = str(value)
        text = text.replace("\r", " ")
        text = text.replace("\u00a0", " ")
        text = text.replace("￾", "-")
        text = re.sub(r"\s+", " ", text)

        return text.strip()

    def _get_extension(self, filename: str) -> str:
        _, extension = os.path.splitext(filename or "")
        return extension.replace(".", "").lower().strip()