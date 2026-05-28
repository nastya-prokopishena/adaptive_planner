import io
import re

import docx
import PyPDF2


class TaskFileExtractorService:
    ALLOWED_EXTENSIONS = {"txt", "pdf", "docx"}

    def extract_text(self, filename, file_bytes):
        extension = filename.rsplit(".", 1)[-1].lower()

        if extension not in self.ALLOWED_EXTENSIONS:
            raise ValueError("Непідтримуваний формат файлу. Доступні: txt, pdf, docx")

        if extension == "txt":
            return self._extract_from_txt(file_bytes)

        if extension == "pdf":
            return self._extract_from_pdf(file_bytes)

        if extension == "docx":
            return self._extract_from_docx(file_bytes)

        return ""

    def _extract_from_txt(self, file_bytes):
        text = file_bytes.decode("utf-8", errors="ignore")
        return self._clean_text(text)

    def _extract_from_pdf(self, file_bytes):
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []

        for page_number, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""

            page_text = self._clean_text(page_text)

            if page_text:
                pages.append(f"\n--- PAGE {page_number} ---\n{page_text}")

        return "\n".join(pages).strip()

    def _extract_from_docx(self, file_bytes):
        document = docx.Document(io.BytesIO(file_bytes))
        parts = []

        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if value:
                parts.append(value)

        for table in document.tables:
            for row in table.rows:
                values = []

                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        values.append(cell_text)

                if values:
                    parts.append(" | ".join(values))

        return self._clean_text("\n".join(parts))

    def _clean_text(self, text):
        text = text or ""
        text = text.replace("\x00", " ")
        text = text.replace("￾", " ")
        text = text.replace("\uf0b7", "•")
        text = text.replace("№ ", "№")

        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s*\n+", "\n", text)

        lines = []

        for line in text.splitlines():
            line = line.strip()

            if not line:
                continue

            if re.fullmatch(r"\d{1,3}", line):
                continue

            lines.append(line)

        return "\n".join(lines).strip()
